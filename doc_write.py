"""
Document-generation tools. Format-aware, not content-aware: they render whatever
structured spec they're handed. Correctness of the spec (does this follow Schedule
III?) is the reference-doc layer's job, not the generator's.

Excel spec shape:
    {"sheets": [
        {"name": "Balance Sheet",
         "rows": [["Particulars", "Note", "31-Mar-25", "31-Mar-24"], ...],
         "number_format_cols": {2: "#,##,##0.00", 3: "#,##,##0.00"},   # Indian grouping
         "bold_rows": [0], "formula_cells": {"C10": "=SUM(C2:C9)"}}
    ]}

Word spec shape:
    {"title": "...", "blocks": [
        {"type": "heading", "text": "...", "level": 1},
        {"type": "para", "text": "..."},
        {"type": "table", "rows": [[...]], "header": True}
    ]}
"""

from __future__ import annotations

import openpyxl
from openpyxl.styles import Font
from docx import Document
from docx.shared import Pt

from .audit_log import audited
from .filesystem import _resolve

# Indian number grouping (lakh/crore): 12,34,56,789.00
INR_FORMAT = "#,##,##0.00"


@audited("generate_excel")
def generate_excel(path: str, spec: dict) -> str:
    p = _resolve(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for sheet in spec["sheets"]:
        ws = wb.create_sheet(title=sheet["name"][:31])
        for r_idx, row in enumerate(sheet.get("rows", []), start=1):
            for c_idx, val in enumerate(row, start=1):
                ws.cell(row=r_idx, column=c_idx, value=val)
        for col, fmt in sheet.get("number_format_cols", {}).items():
            for cell in ws.iter_rows(min_col=int(col) + 1, max_col=int(col) + 1):
                cell[0].number_format = fmt
        for r in sheet.get("bold_rows", []):
            for cell in ws[r + 1]:
                cell.font = Font(bold=True)
        for ref, formula in sheet.get("formula_cells", {}).items():
            ws[ref] = formula
        # auto width
        for col in ws.columns:
            width = max((len(str(c.value)) for c in col if c.value is not None), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(width + 2, 50)
    wb.save(p)
    return str(p)


@audited("generate_word")
def generate_word(path: str, spec: dict) -> str:
    p = _resolve(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if spec.get("title"):
        doc.add_heading(spec["title"], level=0)
    for block in spec.get("blocks", []):
        btype = block["type"]
        if btype == "heading":
            doc.add_heading(block["text"], level=block.get("level", 1))
        elif btype == "para":
            doc.add_paragraph(block["text"])
        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            tbl.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, val in enumerate(row):
                    cell = tbl.cell(r_idx, c_idx)
                    cell.text = str(val)
                    if r_idx == 0 and block.get("header"):
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
    doc.save(p)
    return str(p)


def format_indian_number(value, decimals: int = 2) -> str:
    """Render a number with Indian lakh/crore grouping, e.g. 12345678.5 -> '1,23,45,678.50'."""
    neg = value < 0
    s = f"{abs(value):.{decimals}f}"
    intpart, _, dec = s.partition(".")
    if len(intpart) > 3:
        head, tail = intpart[:-3], intpart[-3:]
        # group the head in pairs from the right
        parts = []
        while len(head) > 2:
            parts.insert(0, head[-2:])
            head = head[:-2]
        if head:
            parts.insert(0, head)
        intpart = ",".join(parts) + "," + tail
    out = f"{intpart}.{dec}" if decimals else intpart
    return f"-{out}" if neg else out
