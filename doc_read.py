"""
Document-reading tools. One router (`ingest_document`) detects the type and
dispatches. Every reader returns a normalized dict so downstream tools don't
care about the source format:

    {"type": ..., "tables": [[...rows...]], "text": "...", "sheets"/"sections": [...],
     "metadata": {...}}

Excel returns BOTH computed values and the formula strings -- CAs care about both.
PDF is intentionally a stub here (text/table extraction is the harder, later piece);
the seam exists so it slots in without touching callers.
"""

from __future__ import annotations

from pathlib import Path

import openpyxl
from docx import Document

from .audit_log import audited
from .filesystem import _resolve


@audited("read_excel", mutating=False)
def read_excel(path: str, sheet: str | None = None) -> dict:
    p = _resolve(path)
    # values workbook (computed) + formulas workbook (raw) so we capture both
    wb_val = openpyxl.load_workbook(p, data_only=True)
    wb_fml = openpyxl.load_workbook(p, data_only=False)
    sheets = []
    names = [sheet] if sheet else wb_val.sheetnames
    for name in names:
        ws_v, ws_f = wb_val[name], wb_fml[name]
        rows, formulas = [], []
        for r_v, r_f in zip(ws_v.iter_rows(values_only=True), ws_f.iter_rows()):
            rows.append(list(r_v))
            formulas.append([c.value if isinstance(c.value, str) and str(c.value).startswith("=")
                             else None for c in r_f])
        sheets.append({"name": name, "rows": rows, "formulas": formulas,
                       "dimensions": ws_v.dimensions})
    return {"type": "xlsx", "sheets": sheets,
            "tables": [s["rows"] for s in sheets],
            "metadata": {"sheet_names": wb_val.sheetnames}}


@audited("read_word", mutating=False)
def read_word(path: str) -> dict:
    p = _resolve(path)
    doc = Document(p)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    tables = []
    for tbl in doc.tables:
        tables.append([[cell.text for cell in row.cells] for row in tbl.rows])
    return {"type": "docx", "text": "\n".join(paragraphs),
            "sections": paragraphs, "tables": tables,
            "metadata": {"n_paragraphs": len(paragraphs), "n_tables": len(tables)}}


@audited("read_csv", mutating=False)
def read_csv(path: str, delimiter: str = ",") -> dict:
    import csv
    p = _resolve(path)
    with open(p, newline="", encoding="utf-8") as fh:
        rows = list(csv.reader(fh, delimiter=delimiter))
    return {"type": "csv", "tables": [rows], "rows": rows,
            "metadata": {"n_rows": len(rows)}}


@audited("read_pdf", mutating=False)
def read_pdf(path: str) -> dict:
    """STUB. Wire pdfplumber (text+tables) with a pytesseract OCR fallback later.
    Kept here so the router and callers are stable before the hard part is built."""
    raise NotImplementedError(
        "PDF reading not yet implemented. Planned: pdfplumber for text/tables, "
        "pytesseract OCR fallback for scanned documents."
    )


_DISPATCH = {".xlsx": read_excel, ".xlsm": read_excel,
             ".docx": read_word, ".csv": read_csv, ".pdf": read_pdf}


@audited("ingest_document", mutating=False)
def ingest_document(path: str) -> dict:
    """Detect file type and dispatch to the right reader. Single entry point."""
    ext = Path(path).suffix.lower()
    reader = _DISPATCH.get(ext)
    if reader is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {sorted(_DISPATCH)}")
    return reader(path)
